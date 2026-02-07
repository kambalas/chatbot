from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument


def extract_text(path: Path) -> str:
    doc = DocxDocument(str(path))
    parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return "\n".join(parts)
